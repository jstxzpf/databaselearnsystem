from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from datetime import datetime

def add_heading_with_style(doc, text, level):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 102, 204)
    return heading

def generate_report():
    document = Document()

    # Style configuration for Chinese support
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # --- Title Page ---
    title = document.add_heading('通用智能学习系统', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 102, 204)
        run.font.size = Pt(28)
    
    subtitle = document.add_paragraph('系统功能报告')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    
    document.add_paragraph()
    
    audience = document.add_paragraph('目标用户：统计调查员')
    audience.alignment = WD_ALIGN_PARAGRAPH.CENTER
    audience.runs[0].font.size = Pt(14)
    
    date_p = document.add_paragraph(f'生成日期：{datetime.now().strftime("%Y年%m月%d日")}')
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_page_break()

    # --- Table of Contents Placeholder ---
    add_heading_with_style(document, '目录', level=1)
    document.add_paragraph('1. 系统简介')
    document.add_paragraph('2. 系统架构')
    document.add_paragraph('3. 核心功能介绍')
    document.add_paragraph('   3.1 智能学习')
    document.add_paragraph('   3.2 智能考试')
    document.add_paragraph('   3.3 智能审批')
    document.add_paragraph('4. 系统特点')
    document.add_paragraph('5. 应用场景')
    document.add_paragraph('6. 使用说明')
    document.add_paragraph('7. 技术架构')
    document.add_paragraph('8. 总结与展望')
    document.add_page_break()

    # --- Introduction ---
    add_heading_with_style(document, '1. 系统简介', level=1)
    
    document.add_heading('1.1 背景', level=2)
    p = document.add_paragraph(
        '随着农业现代化的推进，统计调查工作面临着知识更新快、业务要求高、培训成本大等挑战。'
        '传统的纸质手册和集中培训方式已无法满足分散作业、时效性强的调查工作需求。'
    )
    
    document.add_heading('1.2 系统定位', level=2)
    p = document.add_paragraph(
        '通用智能学习系统是专为农业统计调查员设计的智能化学习辅助工具。'
        '系统以《农产量调查手册》等专业教材为知识库，结合最新的人工智能技术，'
        '提供"即学即练即评"的一体化学习体验，有效提升调查员的业务能力和工作效率。'
    )

    document.add_heading('1.3 核心价值', level=2)
    document.add_paragraph('• 降低培训成本：无需集中培训，随时随地在线学习', style='List Bullet')
    document.add_paragraph('• 提升学习效率：AI智能讲解，个性化学习路径', style='List Bullet')
    document.add_paragraph('• 保证调查质量：考前自测，确保掌握关键知识点', style='List Bullet')
    document.add_paragraph('• 快速能力评估：智能批改，即时反馈学习成果', style='List Bullet')

    # Insert Home screenshot
    if os.path.exists('report_assets/screenshot_home.png'):
        document.add_picture('report_assets/screenshot_home.png', width=Inches(6))
        caption = document.add_paragraph('图 1：系统主界面', style='Caption')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()

    # --- System Architecture ---
    add_heading_with_style(document, '2. 系统架构', level=1)
    
    p = document.add_paragraph(
        '本系统采用现代化的B/S架构设计，由前端展示层、业务逻辑层、AI服务层和数据持久层组成。'
        '前端基于Bootstrap 5框架，响应式设计，支持PC和移动设备访问。'
        '后端使用Python Flask框架，轻量高效。AI服务集成本地Ollama模型，保证数据安全和响应速度。'
    )

    # Insert Architecture Diagram
    if os.path.exists('report_assets/flowchart_architecture.png'):
        document.add_picture('report_assets/flowchart_architecture.png', width=Inches(5.5))
        caption = document.add_paragraph('图 2：系统架构示意图', style='Caption')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()

    # --- System Features ---
    add_heading_with_style(document, '3. 核心功能介绍', level=1)

    # Feature 1: Smart Learning
    add_heading_with_style(document, '3.1 智能学习 (Smart Learning)', level=2)
    
    document.add_heading('功能概述', level=3)
    document.add_paragraph(
        '智能学习模块是系统的核心功能之一，针对《农产量调查手册》等专业资料提供深度智能解析。'
        '调查员可以按章节浏览知识体系，点击任意概念即可获得AI即时生成的详细讲解，'
        '包括定义、应用场景、操作要点和注意事项等。'
    )
    
    document.add_heading('功能特性', level=3)
    document.add_paragraph('• 章节导航：清晰的知识结构，快速定位学习内容', style='List Bullet')
    document.add_paragraph('• AI讲解：基于大语言模型，生成通俗易懂的专业解释', style='List Bullet')
    document.add_paragraph('• 学习进度：自动记录学习轨迹，支持断点续学', style='List Bullet')
    document.add_paragraph('• 知识搜索：关键词快速检索，精准定位知识点', style='List Bullet')
    document.add_paragraph('• 批量生成：一键生成整章讲解，离线学习更方便', style='List Bullet')
    
    document.add_heading('应用示例', level=3)
    document.add_paragraph(
        '例如，调查员在学习"第八章 实割实测调查方法"时，对"样本选择策略"概念不够理解。'
        '点击该知识点后，AI会生成包含理论依据、选择原则、典型案例的详细讲解，'
        '帮助调查员快速掌握科学选点的方法。'
    )

    if os.path.exists('report_assets/screenshot_learning.png'):
        document.add_picture('report_assets/screenshot_learning.png', width=Inches(6))
        caption = document.add_paragraph('图 3：智能学习界面（农产量调查手册）', style='Caption')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()

    # Feature 2: Smart Exam
    add_heading_with_style(document, '3.2 智能考试 (Smart Exam)', level=2)
    
    document.add_heading('功能概述', level=3)
    document.add_paragraph(
        '智能考试模块提供自动化试卷生成功能，支持按章节、题型自定义生成考卷。'
        'AI能够根据知识点自动生成高质量的选择题、判断题、简答题等多种题型，'
        '帮助调查员进行考前自测，检验学习效果。'
    )
    
    document.add_heading('功能特性', level=3)
    document.add_paragraph('• 灵活组卷：支持单章或多章组合，题型自由选择', style='List Bullet')
    document.add_paragraph('• AI出题：题目随机生成，每次考试内容不同', style='List Bullet')
    document.add_paragraph('• 试卷下载：支持TXT格式下载，方便打印和存档', style='List Bullet')
    document.add_paragraph('• 历史记录：自动保存考试历史，随时回顾', style='List Bullet')
    document.add_paragraph('• 题目质量：基于专业知识库，确保题目准确性', style='List Bullet')
    
    document.add_heading('应用示例', level=3)
    document.add_paragraph(
        '调查季开始前，调查员可以选择"数据处理基础"、"数据审核流程"等核心章节，'
        '生成一套包含20道选择题和5道简答题的综合测试卷，进行考前强化复习。'
    )

    if os.path.exists('report_assets/screenshot_exam.png'):
        document.add_picture('report_assets/screenshot_exam.png', width=Inches(6))
        caption = document.add_paragraph('图 4：智能考试生成界面', style='Caption')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()

    # Feature 3: Smart Review
    add_heading_with_style(document, '3.3 智能审批 (Smart Review)', level=2)
    
    document.add_heading('功能概述', level=3)
    document.add_paragraph(
        '智能审批模块利用AI技术对上传的试卷或调查报告进行自动批改和分析。'
        '系统能够识别答题内容，指出错误点，生成详细的改进建议和薄弱知识点分析，'
        '大幅提升培训和考核效率。'
    )
    
    document.add_heading('功能特性', level=3)
    document.add_paragraph('• 文件上传：支持TXT、PDF、DOC等多种格式', style='List Bullet')
    document.add_paragraph('• 智能批改：AI自动分析答题情况，给出评分和建议', style='List Bullet')
    document.add_paragraph('• 错误定位：精准指出知识盲点和常见错误', style='List Bullet')
    document.add_paragraph('• 学习建议：基于批改结果，推荐针对性学习内容', style='List Bullet')
    document.add_paragraph('• 审批历史：完整记录所有批改结果，便于对比进步', style='List Bullet')
    
    document.add_heading('应用示例', level=3)
    document.add_paragraph(
        '新入职调查员完成培训测试后，将答卷上传至系统。AI自动批改后指出：'
        '"对实割实测样本选择理解不足，建议重点学习第八章相关内容"，'
        '帮助培训主管快速掌握新员工能力水平。'
    )

    if os.path.exists('report_assets/screenshot_review.png'):
        document.add_picture('report_assets/screenshot_review.png', width=Inches(6))
        caption = document.add_paragraph('图 5：智能审批/上传界面', style='Caption')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()

    # --- System Features ---
    add_heading_with_style(document, '4. 系统特点', level=1)
    
    document.add_heading('4.1 人工智能驱动', level=2)
    p = document.add_paragraph(
        '系统基于先进的大语言模型（如Qwen、Gemma等），具备强大的自然语言理解和生成能力。'
        'AI不仅能准确理解专业术语，还能用通俗易懂的语言进行讲解，降低学习门槛。'
    )
    
    document.add_heading('4.2 数据安全可靠', level=2)
    p = document.add_paragraph(
        '采用本地部署的Ollama AI服务，所有数据处理在本地完成，不依赖外部云服务，'
        '有效保护调查数据的机密性和安全性，符合统计工作保密要求。'
    )
    
    document.add_heading('4.3 操作简单便捷', level=2)
    p = document.add_paragraph(
        '界面设计遵循现代UI/UX原则，操作逻辑清晰直观。调查员无需专业培训，'
        '即可快速上手使用，降低了系统推广难度。'
    )
    
    document.add_heading('4.4 跨平台兼容', level=2)
    p = document.add_paragraph(
        '基于Web技术开发，支持Windows、Linux等多种操作系统，'
        '兼容PC、平板、手机等不同终端，真正实现"随时随地学习"。'
    )
    
    document.add_heading('4.5 扩展性强', level=2)
    p = document.add_paragraph(
        '采用模块化设计，知识库可灵活扩展。除农产量调查手册外，'
        '系统已支持Linux系统操作、统计执法证考试等多门课程，未来可根据需求快速添加新课程。'
    )

    document.add_page_break()

    # --- Usage Scenarios ---
    add_heading_with_style(document, '5. 应用场景', level=1)
    
    document.add_heading('场景一：调查前强化复习', level=2)
    document.add_paragraph('【背景】调查季开始前，调查员需要温习调查方法和流程')
    document.add_paragraph(
        '【解决方案】调查员登录系统，选择"调查基础知识"、"实割实测调查方法"等核心章节进行复习。'
        '对于遗忘的概念，通过AI讲解功能快速回忆。复习完成后，生成一套综合测试卷进行自测，'
        '确保掌握关键知识点后再开展实地调查工作。'
    )
    
    document.add_heading('场景二：新进人员快速培训', level=2)
    document.add_paragraph('【背景】新入职统计人员需要快速掌握业务知识')
    document.add_paragraph(
        '【解决方案】培训主管为新员工指定学习路径，包括"调查基础知识"、"法律法规与伦理规范"等必修章节。'
        '新员工按章节自主学习，系统自动记录学习进度。培训结束后，主管生成标准化考卷，'
        '新员工完成后上传系统，AI自动批改并生成能力评估报告，主管根据报告安排后续培训。'
    )
    
    document.add_heading('场景三：日常知识查询', level=2)
    document.add_paragraph('【背景】调查员在工作中遇到疑问，需要快速查阅')
    document.add_paragraph(
        '【解决方案】调查员使用搜索功能输入关键词，如"数据清洗"，系统快速定位到相关章节和知识点，'
        '并提供AI详细讲解。相比传统查阅纸质手册，效率提升数倍。'
    )
    
    document.add_heading('场景四：能力持续提升', level=2)
    document.add_paragraph('【背景】资深调查员需要不断更新知识，保持专业能力')
    document.add_paragraph(
        '【解决方案】定期使用智能考试功能进行自测，检验对新政策、新方法的掌握情况。'
        '系统会根据答题情况推荐薄弱知识点，实现针对性学习，持续提升业务能力。'
    )

    # Insert Usage Flowchart
    if os.path.exists('report_assets/flowchart_usage.png'):
        document.add_picture('report_assets/flowchart_usage.png', width=Inches(5.5))
        caption = document.add_paragraph('图 6：典型使用流程', style='Caption')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()

    # --- User Manual ---
    add_heading_with_style(document, '6. 使用说明', level=1)
    
    document.add_heading('6.1 访问系统', level=2)
    document.add_paragraph('1. 打开浏览器（推荐使用Chrome、Edge或Firefox）')
    document.add_paragraph('2. 在地址栏输入系统地址（如：http://localhost:5000）')
    document.add_paragraph('3. 系统自动加载主界面，无需登录即可使用')
    
    document.add_heading('6.2 切换课程', level=2)
    document.add_paragraph('1. 点击页面右上角的"设置"按钮')
    document.add_paragraph('2. 在课程下拉列表中选择需要学习的课程（如"农产量调查手册"）')
    document.add_paragraph('3. 系统自动切换到选定课程，返回主页即可开始学习')
    
    document.add_heading('6.3 智能学习操作', level=2)
    document.add_paragraph('1. 点击导航栏的"学习"菜单进入学习界面')
    document.add_paragraph('2. 左侧显示章节列表，点击章节名称展开该章内容')
    document.add_paragraph('3. 中间区域显示主要概念和知识点')
    document.add_paragraph('4. 点击任意概念，右侧自动显示AI生成的详细讲解')
    document.add_paragraph('5. 如需重新生成讲解，可点击"重新生成"按钮')
    document.add_paragraph('6. 使用搜索框可快速查找特定知识点')
    
    document.add_heading('6.4 智能考试操作', level=2)
    document.add_paragraph('1. 点击导航栏的"考试"菜单进入考试界面')
    document.add_paragraph('2. 左侧勾选需要考查的章节（可多选）')
    document.add_paragraph('3. 右侧选择题型（如选择题、简答题等）')
    document.add_paragraph('4. 开启"AI生成"开关可获得更高质量的题目')
    document.add_paragraph('5. 点击"生成试卷"按钮，等待片刻后试卷自动显示')
    document.add_paragraph('6. 点击"下载试卷"可将试卷保存为文本文件')
    
    document.add_heading('6.5 智能审批操作', level=2)
    document.add_paragraph('1. 点击导航栏的"审批"菜单进入审批界面')
    document.add_paragraph('2. 点击"选择文件"按钮，选择要批改的试卷文件')
    document.add_paragraph('3. 文件上传后，系统自动预览试卷内容')
    document.add_paragraph('4. 点击"开始审批"按钮，AI开始分析答题情况')
    document.add_paragraph('5. 审批完成后，查看详细的批改结果和学习建议')
    document.add_paragraph('6. 在"审批历史"中可查看所有批改记录')

    document.add_page_break()

    # --- Technical Architecture ---
    add_heading_with_style(document, '7. 技术架构', level=1)
    
    document.add_heading('7.1 技术选型', level=2)
    document.add_paragraph('• 后端框架：Flask 2.3.3（轻量级Python Web框架）', style='List Bullet')
    document.add_paragraph('• 前端技术：Bootstrap 5 + jQuery（响应式设计）', style='List Bullet')
    document.add_paragraph('• 数据库：SQLite（轻量级关系型数据库）', style='List Bullet')
    document.add_paragraph('• AI服务：Ollama本地API（支持多种开源大模型）', style='List Bullet')
    document.add_paragraph('• 数据格式：JSON（知识库存储）', style='List Bullet')
    
    document.add_heading('7.2 系统优势', level=2)
    document.add_paragraph('• 部署简单：单机即可运行，无需复杂环境配置', style='List Bullet')
    document.add_paragraph('• 资源占用低：适合在普通办公电脑上运行', style='List Bullet')
    document.add_paragraph('• 维护成本低：开源技术栈，无授权费用', style='List Bullet')
    document.add_paragraph('• 可定制性强：源代码开放，可根据需求二次开发', style='List Bullet')

    document.add_page_break()

    # --- Conclusion ---
    add_heading_with_style(document, '8. 总结与展望', level=1)
    
    document.add_heading('8.1 系统价值总结', level=2)
    p = document.add_paragraph(
        '通用智能学习系统通过引入人工智能技术，有效解决了传统培训模式中存在的"学练脱节"、'
        '"培训成本高"、"效果难评估"等痛点问题。系统不仅是一个学习平台，'
        '更是统计调查员随身携带的智能助手，随时随地提供专业支持。'
    )
    
    document.add_heading('8.2 未来展望', level=2)
    document.add_paragraph(
        '未来，系统将朝着以下方向持续优化：'
    )
    document.add_paragraph('• 知识库扩展：增加更多专业课程，覆盖统计工作全领域', style='List Bullet')
    document.add_paragraph('• 学习路径优化：基于学习数据，推荐个性化学习计划', style='List Bullet')
    document.add_paragraph('• 移动应用开发：推出独立移动App，提升移动端体验', style='List Bullet')
    document.add_paragraph('• 团队协作功能：支持团队学习、成绩统计等协作功能', style='List Bullet')
    document.add_paragraph('• 多模态交互：引入语音问答、图像识别等新型交互方式', style='List Bullet')
    
    document.add_paragraph()
    closing = document.add_paragraph('我们相信，随着技术的不断进步，智能学习系统必将在统计调查工作中发挥越来越重要的作用！')
    closing.runs[0].font.bold = True

    # Save document
    document.save('system_function_report.docx')
    print("Enhanced report generated: system_function_report.docx")

if __name__ == '__main__':
    generate_report()
